from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from ebooklib import ITEM_DOCUMENT, epub
from pypdf import PdfReader

from .types import ExtractedDocument, ExtractedSegment, ManifestSource

logger = logging.getLogger(__name__)


def _normalize_text(value: str) -> str:
    lines = [line.strip() for line in value.splitlines()]
    filtered = [line for line in lines if line]
    return "\n".join(filtered).strip()


def _extract_pdf(path: Path, source: ManifestSource) -> ExtractedDocument | None:
    try:
        reader = PdfReader(str(path))
    except Exception:
        logger.warning("Failed to read PDF %s", path.name, exc_info=True)
        return None

    segments: list[ExtractedSegment] = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = _normalize_text(page.extract_text() or "")
        if not text:
            continue
        segments.append(ExtractedSegment(text=text, page_number=page_number))

    if not segments:
        logger.warning("No extractable PDF text found in %s", path.name)
        return None

    return ExtractedDocument(source=source, path=path, segments=segments)


def _extract_docx(path: Path, source: ManifestSource) -> ExtractedDocument | None:
    try:
        document = DocxDocument(str(path))
    except Exception:
        logger.warning("Failed to read DOCX %s", path.name, exc_info=True)
        return None

    segments: list[ExtractedSegment] = []
    current_section: str | None = None

    for paragraph in document.paragraphs:
        text = _normalize_text(paragraph.text)
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            current_section = text
            continue
        segments.append(ExtractedSegment(text=text, section=current_section))

    if not segments:
        logger.warning("No extractable DOCX text found in %s", path.name)
        return None

    return ExtractedDocument(source=source, path=path, segments=segments)


@dataclass(slots=True, frozen=True)
class _HtmlResult:
    segments: list[ExtractedSegment]
    title: str | None = None


def _extract_structured_html(html: str) -> _HtmlResult:
    soup = BeautifulSoup(html, "html.parser")
    container = soup.body or soup
    current_section: str | None = None
    segments: list[ExtractedSegment] = []

    title_element = soup.find(["h1", "h2", "title"])
    title = _normalize_text(title_element.get_text(" ", strip=True)) if title_element else None

    for element in container.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
        text = _normalize_text(element.get_text(" ", strip=True))
        if not text:
            continue
        if element.name and element.name.startswith("h"):
            current_section = text
            continue
        segments.append(ExtractedSegment(text=text, section=current_section))

    if segments:
        return _HtmlResult(segments=segments, title=title)

    fallback_text = _normalize_text(container.get_text("\n", strip=True))
    if not fallback_text:
        return _HtmlResult(segments=[], title=title)
    return _HtmlResult(segments=[ExtractedSegment(text=fallback_text)], title=title)


def _extract_epub(path: Path, source: ManifestSource) -> ExtractedDocument | None:
    try:
        book = epub.read_epub(str(path))
    except Exception:
        logger.warning("Failed to read EPUB %s", path.name, exc_info=True)
        return None

    segments: list[ExtractedSegment] = []

    for item in book.get_items():
        if item.get_type() != ITEM_DOCUMENT:
            continue
        item_name = item.get_name().lower()
        if item_name.startswith("nav") or "toc" in item_name:
            continue

        body_content = item.get_body_content()
        if body_content is None:
            continue
        html = body_content.decode("utf-8", errors="ignore")
        result = _extract_structured_html(html)
        if not result.segments:
            continue

        chapter_title = result.title
        for segment in result.segments:
            section = segment.section or chapter_title or item.get_name()
            segments.append(
                ExtractedSegment(
                    text=segment.text,
                    section=section,
                    page_number=segment.page_number,
                )
            )

    if not segments:
        logger.warning("No extractable EPUB text found in %s", path.name)
        return None

    return ExtractedDocument(source=source, path=path, segments=segments)


def _extract_html(path: Path, source: ManifestSource) -> ExtractedDocument | None:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        logger.warning("Failed to decode HTML %s as UTF-8, retrying with replace", path.name)
        text = path.read_text(encoding="utf-8", errors="replace")

    result = _extract_structured_html(text)
    if not result.segments:
        logger.warning("No extractable HTML text found in %s", path.name)
        return None
    return ExtractedDocument(source=source, path=path, segments=result.segments)


def extract_document(path: Path, source: ManifestSource) -> ExtractedDocument | None:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path, source)
    if suffix == ".epub":
        return _extract_epub(path, source)
    if suffix == ".docx":
        return _extract_docx(path, source)
    if suffix in {".html", ".htm"}:
        return _extract_html(path, source)

    logger.warning("Unsupported file type for %s", path.name)
    return None
