from __future__ import annotations

from pathlib import Path

from docx import Document
from ebooklib import epub
from src.knowledge_ingestion.types import ManifestSource


def make_manifest_source(**overrides: object) -> ManifestSource:
    values: dict[str, object] = {
        "file": "fixture.pdf",
        "source_type": "article",
        "title": "Fixture Title",
        "language": "en",
        "author": "Fixture Author",
        "url": None,
        "tags": ["fixture"],
    }
    values.update(overrides)
    return ManifestSource(**values)


def _pdf_stream(text: str) -> str:
    escaped = (
        text.replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
        .replace("\n", "\\n")
    )
    return f"BT /F1 18 Tf 50 100 Td ({escaped}) Tj ET"


def write_pdf(path: Path, pages: list[str]) -> None:
    objects: list[str] = []
    kids = [f"{3 + index} 0 R" for index in range(len(pages))]
    objects.append("<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(f"<< /Type /Pages /Kids [{' '.join(kids)}] /Count {len(pages)} >>")

    next_object_id = 3 + len(pages)
    font_id = next_object_id + len(pages)
    for index, _ in enumerate(pages):
        content_id = next_object_id + index
        objects.append(
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] "
            f"/Contents {content_id} 0 R /Resources << /Font << /F1 {font_id} 0 R >> >> >>"
        )

    for page_text in pages:
        stream = _pdf_stream(page_text)
        stream_bytes_len = len(stream.encode("utf-8"))
        objects.append(f"<< /Length {stream_bytes_len} >>\nstream\n{stream}\nendstream")

    objects.append("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    lines = ["%PDF-1.4\n"]
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(sum(len(line.encode("utf-8")) for line in lines))
        lines.append(f"{index} 0 obj\n{obj}\nendobj\n")

    xref_offset = sum(len(line.encode("utf-8")) for line in lines)
    lines.append(f"xref\n0 {len(objects) + 1}\n")
    lines.append("0000000000 65535 f \n")
    for offset in offsets[1:]:
        lines.append(f"{offset:010} 00000 n \n")
    lines.append(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n"
    )

    path.write_bytes("".join(lines).encode("utf-8"))


def write_docx(path: Path, heading: str, paragraphs: list[str]) -> None:
    document = Document()
    document.add_heading(heading, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def write_epub(path: Path, chapters: list[tuple[str, list[str]]]) -> None:
    book = epub.EpubBook()
    book.set_identifier("fixture-book")
    book.set_title("Fixture EPUB")
    book.set_language("en")

    items = []
    for index, (title, paragraphs) in enumerate(chapters, start=1):
        chapter = epub.EpubHtml(
            title=title,
            file_name=f"chapter-{index}.xhtml",
            lang="en",
        )
        body = "".join(f"<p>{paragraph}</p>" for paragraph in paragraphs)
        chapter.content = f"<html><body><h1>{title}</h1>{body}</body></html>"
        book.add_item(chapter)
        items.append(chapter)

    book.toc = tuple(items)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", *items]
    epub.write_epub(str(path), book)
